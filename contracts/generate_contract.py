#!/usr/bin/env python3
"""
Contract Generator — Tạo hợp đồng DOCX từ template Markdown.
Sử dụng: python3 generate_contract.py --template <path> --output <path> --data '<json>'
"""

import argparse
import json
import sys
import re
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, Cm, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
except ImportError:
    print("❌ Cần cài python-docx: pip3 install python-docx")
    sys.exit(1)


def fill_template(template_text: str, data: dict) -> str:
    """Replace {{VARIABLE}} placeholders with actual values."""
    result = template_text
    for key, value in data.items():
        result = result.replace(f"{{{{{key}}}}}", str(value))
    
    # Warn about unfilled placeholders
    remaining = re.findall(r'\{\{(\w+)\}\}', result)
    if remaining:
        print(f"⚠️ Các trường chưa điền: {', '.join(remaining)}")
    
    return result


def markdown_to_docx(md_text: str, output_path: str):
    """Convert filled markdown template to DOCX."""
    doc = Document()
    
    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin = Cm(3)
        section.right_margin = Cm(2)
    
    lines = md_text.strip().split('\n')
    
    for line in lines:
        stripped = line.strip()
        
        if not stripped:
            doc.add_paragraph('')
            continue
        
        # Headers
        if stripped.startswith('# '):
            p = doc.add_heading(stripped[2:], level=0)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        elif stripped.startswith('## '):
            p = doc.add_heading(stripped[3:], level=1)
        elif stripped.startswith('### '):
            p = doc.add_heading(stripped[4:], level=2)
        elif stripped.startswith('---'):
            # Horizontal rule → thin paragraph
            p = doc.add_paragraph('─' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            p.runs[0].font.size = Pt(8)
            p.runs[0].font.color.rgb = RGBColor(150, 150, 150)
        elif stripped.startswith('- ') or stripped.startswith('* '):
            # Bullet point
            doc.add_paragraph(stripped[2:], style='List Bullet')
        else:
            # Regular paragraph with bold support
            p = doc.add_paragraph()
            # Handle **bold** markers
            parts = re.split(r'(\*\*[^*]+\*\*)', stripped)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.size = Pt(12)
                else:
                    run = p.add_run(part)
                    run.font.size = Pt(12)
    
    # Set default font
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    doc.save(output_path)
    print(f"✅ Đã tạo file: {output_path}")


def main():
    parser = argparse.ArgumentParser(description='Tạo hợp đồng từ template')
    parser.add_argument('--template', required=True, help='Path tới file template (.md)')
    parser.add_argument('--output', required=True, help='Path output file (.docx)')
    parser.add_argument('--data', required=True, help='JSON string chứa dữ liệu điền vào')
    
    args = parser.parse_args()
    
    # Read template
    template_path = Path(args.template)
    if not template_path.exists():
        print(f"❌ Template không tồn tại: {args.template}")
        sys.exit(1)
    
    template_text = template_path.read_text(encoding='utf-8')
    
    # Parse data
    try:
        data = json.loads(args.data)
    except json.JSONDecodeError as e:
        print(f"❌ Lỗi JSON data: {e}")
        sys.exit(1)
    
    # Fill template
    filled = fill_template(template_text, data)
    
    # Create output directory
    output_path = Path(args.output)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    
    # Generate DOCX
    markdown_to_docx(filled, str(output_path))
    
    print(f"📄 Hợp đồng tạo thành công!")
    print(f"📁 File: {output_path}")
    print(f"📊 Kích thước: {output_path.stat().st_size / 1024:.1f} KB")


if __name__ == '__main__':
    main()
