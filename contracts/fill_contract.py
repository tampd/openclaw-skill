#!/usr/bin/env python3
"""
BKNS Contract Filler — Thay thế thông tin khách hàng trực tiếp trong file DOCX.
Giữ nguyên toàn bộ format, font, table, layout gốc.

Sử dụng:
  python3 fill_contract.py --template <path.docx> --output <path.docx> --data '<json>'

Ví dụ:
  python3 fill_contract.py \
    --template /root/docs/HD_Temp/HD_hosting.docx \
    --output /root/openclaw/workspace/contracts/output/HD-ABC-hosting.docx \
    --data '{"KHACH_TEN": "Công ty ABC", "KHACH_DIACHI": "123 Lê Lợi Q1", ...}'
"""

import argparse
import json
import sys
import copy
from pathlib import Path

try:
    from docx import Document
except ImportError:
    print("❌ Cần cài python-docx: pip3 install --break-system-packages python-docx")
    sys.exit(1)


def replace_text_in_paragraph(paragraph, old_text, new_text):
    """
    Replace text in a paragraph while preserving formatting.
    Handles text split across multiple runs.
    """
    # Method 1: Try simple run-level replacement first
    for run in paragraph.runs:
        if old_text in run.text:
            run.text = run.text.replace(old_text, new_text)
            return True
    
    # Method 2: Full paragraph text replacement (for text split across runs)
    full_text = paragraph.text
    if old_text in full_text:
        new_full = full_text.replace(old_text, new_text)
        # Clear all runs and set text on the first one (preserving its formatting)
        if paragraph.runs:
            first_run = paragraph.runs[0]
            first_run_format = copy.deepcopy(first_run.font)
            for run in paragraph.runs:
                run.text = ""
            first_run.text = new_full
        return True
    
    return False


def replace_text_in_table(table, old_text, new_text):
    """Replace text in all cells of a table."""
    replaced = False
    for row in table.rows:
        for cell in row.cells:
            for paragraph in cell.paragraphs:
                if replace_text_in_paragraph(paragraph, old_text, new_text):
                    replaced = True
    return replaced


def fill_contract(template_path, output_path, replacements):
    """
    Copy template DOCX and replace specified text.
    Preserves all formatting, tables, headers, footers.
    """
    doc = Document(template_path)
    
    results = {}
    
    for old_text, new_text in replacements.items():
        found = False
        
        # Replace in paragraphs
        for para in doc.paragraphs:
            if replace_text_in_paragraph(para, old_text, new_text):
                found = True
        
        # Replace in tables
        for table in doc.tables:
            if replace_text_in_table(table, old_text, new_text):
                found = True
        
        # Replace in headers/footers
        for section in doc.sections:
            if section.header:
                for para in section.header.paragraphs:
                    if replace_text_in_paragraph(para, old_text, new_text):
                        found = True
            if section.footer:
                for para in section.footer.paragraphs:
                    if replace_text_in_paragraph(para, old_text, new_text):
                        found = True
        
        results[old_text] = "✅ Đã thay" if found else "⚠️ Không tìm thấy"
    
    # Save
    output_path = Path(output_path)
    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    
    return results


def main():
    parser = argparse.ArgumentParser(description='Điền thông tin vào hợp đồng BKNS')
    parser.add_argument('--template', required=True, help='Path tới file DOCX template')
    parser.add_argument('--output', required=True, help='Path output file DOCX')
    parser.add_argument('--data', required=True, help='JSON: {"text cũ": "text mới", ...}')
    
    args = parser.parse_args()
    
    template_path = Path(args.template)
    if not template_path.exists():
        print(f"❌ Template không tồn tại: {args.template}")
        sys.exit(1)
    
    try:
        replacements = json.loads(args.data)
    except json.JSONDecodeError as e:
        print(f"❌ Lỗi JSON: {e}")
        sys.exit(1)
    
    print(f"📄 Template: {template_path.name}")
    print(f"📝 Thay thế {len(replacements)} trường...")
    
    results = fill_contract(str(template_path), args.output, replacements)
    
    print(f"\n📊 Kết quả:")
    for old_text, status in results.items():
        display = old_text[:50] + "..." if len(old_text) > 50 else old_text
        print(f"  {status} | '{display}'")
    
    output_path = Path(args.output)
    print(f"\n✅ File đã tạo: {output_path}")
    print(f"📊 Kích thước: {output_path.stat().st_size / 1024:.1f} KB")


if __name__ == '__main__':
    main()
